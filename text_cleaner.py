#!/usr/bin/env python3
"""
Скрипт для очистки текста от цифровых водяных знаков ИИ.
Удаляет невидимые символы Unicode из текстовых файлов.
"""

import sys
import os
import argparse
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Словарь замен: {исходный_символ: символ_замены}
UNICODE_REPLACEMENTS = {
    '\u00A0': '\u0020',  # Неразрывный пробел -> Обычный пробел
    '\u202F': '\u0020',  # Узкий неразрывный пробел -> Обычный пробел
    '\u200B': '',        # Пробел нулевой ширины -> Удалить
    '\u00AD': '',        # Мягкий перенос -> Удалить
    '\u2019': '\u0027',  # Правый одинарный апостроф -> Обычный апостроф
}


def clean_text(text):
    """
    Очищает текст от невидимых символов Unicode.

    Args:
        text (str): Исходный текст

    Returns:
        str: Очищенный текст
    """
    cleaned_text = text
    for old_char, new_char in UNICODE_REPLACEMENTS.items():
        cleaned_text = cleaned_text.replace(old_char, new_char)
    return cleaned_text


def clean_txt_file(input_path, output_path):
    """
    Очищает текстовый файл (.txt).

    Args:
        input_path (str): Путь к исходному файлу
        output_path (str): Путь к выходному файлу
    """
    try:
        # Чтение файла с автоопределением кодировки
        encodings = ['utf-8', 'cp1251', 'windows-1251', 'latin-1']
        text = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(input_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    used_encoding = encoding
                    break
            except UnicodeDecodeError:
                continue

        if text is None:
            raise ValueError(f"Не удалось прочитать файл с известными кодировками")

        # Очистка текста
        cleaned_text = clean_text(text)

        # Сохранение результата
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(cleaned_text)

        print(f"✓ Текстовый файл успешно очищен")
        print(f"  Исходная кодировка: {used_encoding}")
        print(f"  Результат сохранён в: {output_path}")

    except Exception as e:
        print(f"✗ Ошибка при обработке текстового файла: {e}", file=sys.stderr)
        sys.exit(1)


def clean_docx_file(input_path, output_path):
    """
    Очищает документ Word (.docx).

    Args:
        input_path (str): Путь к исходному файлу
        output_path (str): Путь к выходному файлу
    """
    if not DOCX_AVAILABLE:
        print("✗ Библиотека python-docx не установлена!", file=sys.stderr)
        print("  Установите её командой: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    try:
        # Открытие документа
        doc = Document(input_path)

        # Очистка параграфов
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = clean_text(run.text)

        # Очистка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = clean_text(run.text)

        # Сохранение документа
        doc.save(output_path)

        print(f"✓ Документ Word успешно очищен")
        print(f"  Результат сохранён в: {output_path}")

    except Exception as e:
        print(f"✗ Ошибка при обработке документа Word: {e}", file=sys.stderr)
        sys.exit(1)


def get_file_extension(file_path):
    """Возвращает расширение файла в нижнем регистре."""
    return Path(file_path).suffix.lower()


def generate_output_path(input_path, overwrite=False):
    """
    Генерирует путь для выходного файла.

    Args:
        input_path (str): Путь к исходному файлу
        overwrite (bool): Если True, перезаписывает исходный файл

    Returns:
        str: Путь к выходному файлу
    """
    if overwrite:
        return input_path

    path = Path(input_path)
    stem = path.stem
    suffix = path.suffix
    parent = path.parent

    return str(parent / f"{stem}_cleaned{suffix}")


def main():
    """Главная функция скрипта."""
    parser = argparse.ArgumentParser(
        description='Очистка текста от цифровых водяных знаков ИИ (невидимых символов Unicode)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры использования:
  %(prog)s input.txt                    # Создаст input_cleaned.txt
  %(prog)s input.txt -o output.txt      # Сохранит в output.txt
  %(prog)s input.docx --overwrite       # Перезапишет исходный файл
  %(prog)s document.docx -o clean.docx  # Обработает Word документ

Заменяемые символы:
  • Неразрывный пробел (U+00A0) → Обычный пробел
  • Узкий неразрывный пробел (U+202F) → Обычный пробел
  • Пробел нулевой ширины (U+200B) → Удаляется
  • Мягкий перенос (U+00AD) → Удаляется
  • Правый одинарный апостроф (U+2019) → Обычный апостроф (')
        """
    )

    parser.add_argument(
        'input_file',
        help='Путь к исходному файлу (.txt или .docx)'
    )

    parser.add_argument(
        '-o', '--output',
        dest='output_file',
        help='Путь к выходному файлу (по умолчанию: <имя>_cleaned.<расширение>)'
    )

    parser.add_argument(
        '--overwrite',
        action='store_true',
        help='Перезаписать исходный файл вместо создания нового'
    )

    parser.add_argument(
        '--version',
        action='version',
        version='%(prog)s 1.0'
    )

    args = parser.parse_args()

    # Проверка существования входного файла
    if not os.path.exists(args.input_file):
        print(f"✗ Файл не найден: {args.input_file}", file=sys.stderr)
        sys.exit(1)

    # Определение выходного файла
    if args.output_file:
        output_path = args.output_file
    else:
        output_path = generate_output_path(args.input_file, args.overwrite)

    # Получение расширения файла
    file_ext = get_file_extension(args.input_file)

    # Вывод информации
    print(f"\n{'='*60}")
    print(f"Очистка текста от невидимых символов Unicode")
    print(f"{'='*60}")
    print(f"Исходный файл: {args.input_file}")
    print(f"Выходной файл: {output_path}")
    print(f"{'='*60}\n")

    # Обработка файла в зависимости от расширения
    if file_ext == '.txt':
        clean_txt_file(args.input_file, output_path)
    elif file_ext == '.docx':
        clean_docx_file(args.input_file, output_path)
    else:
        print(f"✗ Неподдерживаемый формат файла: {file_ext}", file=sys.stderr)
        print(f"  Поддерживаются только .txt и .docx файлы", file=sys.stderr)
        sys.exit(1)

    print(f"\n{'='*60}")
    print(f"Готово!")
    print(f"{'='*60}\n")


if __name__ == '__main__':
    main()
